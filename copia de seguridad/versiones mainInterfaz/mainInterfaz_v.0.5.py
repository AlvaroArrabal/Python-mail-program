from tkinter import *
from tkinter import messagebox
from modules import SRVCC,CA
import time
from docx import Document

textoPUSCH = "- Valores elevados de PUSCH"
textoRSSI = "- Valores elevados de RSSI"
textoSRVCC = "- No hay intentos de SRVCC"
textoTrafico5G = "- No hay valores de tráfico 5G"
textoMIMORank2Bajo = "- Valores bajos de MIMO Rank2"
textoSinMIMORank2 = "- No hay valores de MIMO Rank2"
textoMIMORank4Bajo = "- Valores bajos de MIMO Rank4"
textoSinMIMORank4 = "- No hay valores de MIMO Rank4"
textoSinCAPCELL = "- No hay valores de CA Pcell"
textoSinCASCELL = "- No hay valores de CA Scell"
word = Document()

def check_button():
    cont=0
    for i in techBox:
        if optionsVariables[cont].get() == 1:
            sectorBox[cont].config(state=NORMAL)
            if sectorBox[cont].get() == "0":
                sectorBox[cont].delete(0,END)
                sectorBox[cont].focus()
            
            techBox[cont].config(state=NORMAL)
            if techBox[cont].get() == "0":
                techBox[cont].delete(0,END)
                techBox[cont].focus()
        else:
            techBox[cont].config(state=DISABLED)
            techText[cont].set("0")
            sectorBox[cont].config(state=DISABLED)
            sectorText[cont].set("0")
        cont += 1

def saveSRVCC():
    tech = techTextSRVCC.get().upper()
    node = nodeTextSRVCC.get()
    cells = numberSectorsTextSRVCC.get()

    try:
        if tech == "0" or node == "0" or cells == "0":
            messagebox.showinfo("Information","Error: Not enough information")
        else:
            SRVCC.crear_txt(tech,node,int(cells))
            messagebox.showinfo("Information","Completed!")
    except:
        messagebox.showinfo("Information","Something went wrong :(")

def saveCA():
    try:
        
        site = siteName.get().upper()
        techCA = techTextCA.get().upper()
        otherTechs = otherTechTextCA.get().upper()
        cells = numberSectorsTextCA.get()
        node = nodeTextCA.get().upper()
        enodeB = enodeTextCA.get().upper()
        if site == "":
            messagebox.showinfo("Information","Error: Empty site name")
        elif techCA == "0" or otherTechs == "0" or cells == "0" or node == "0" or enodeB == "0":
            messagebox.showinfo("Information","Error: Not enough information")
        elif techCA == "Y" and "M" in otherTechs and len(otherTechs) == 1:
            messagebox.showinfo("Information","Error: Not necessary CA between Y and M.")
        else:
            CA.crear_txt_CA(site,techCA,otherTechs,int(cells),node,enodeB)
            messagebox.showinfo("Information","Completed!")
    except:
        messagebox.showinfo("Information","Something went wrong :(")

def write_mail():
    mailText.delete(1.0 ,END)
    now = time.strftime("%X")
    
    if now < str(12):
        mailText.insert(END,"Buenos días,\n")
        word.add_paragraph("Buenos días,")
    else:
        mailText.insert(END,"Buenas tardes,\n")
        word.add_paragraph("Buenas tardes,")

    mailText.insert(END,f"Hemos detectado las siguientes inconsistencias en el site {siteName.get().upper()}:\n")
    word.add_paragraph(f"Hemos detectado las siguientes inconsistencias en el site {siteName.get().upper()}:")
    cont = 0
    for i in techText:
        if i.get() != "0":
            textoAux = ""
            if optionList[cont] == "PUSCH":
                textoAux = ""
                mailText.insert(END,f"{textoPUSCH}")
                textoAux += f"{textoPUSCH}"
                listAux = sectorText[cont].get().split("/")
                techaux = techText[cont].get().upper().split()
                cont2 = 0
                sectorAux = listAux[0].split()
                
                if len(listAux) == 1 and len(sectorAux) == 1:
                    mailText.insert(END," en el sector")
                    textoAux += " en el sector"
                else:
                    mailText.insert(END," en los sectores")
                    textoAux += " en los sectores"
                

                for j in techaux:
                    sectorNumbers = listAux[cont2].split()
                    #print(f"cont2 - {cont2} and {len(techaux)}")
                    
                    cont3 = 0
                    for k in sectorNumbers:
                        mailText.insert(END,f" {j}{k}")
                        textoAux += f" {j}{k}"
                        #print(f"{cont3} - {len(sectorNumbers)} - {sectorNumbers[len(sectorNumbers)-1]} -{k} - {cont2} - {len(techaux)-1}")
                        if cont3 < len(sectorNumbers)   and cont2 != len(techaux)-1:
                            mailText.insert(END,",")
                            textoAux += ","
                        else:
                            if cont3 < len(sectorNumbers)-1:
                                mailText.insert(END,",")
                                textoAux += ","
                            else:
                                mailText.insert(END,".")
                                textoAux += "."
                        cont3 += 1
                    cont2 += 1
                     
                
            if optionList[cont] == "RSSI":
                textoAux = ""
                mailText.insert(END,f"{textoRSSI}")
                textoAux += f"{textoRSSI}"
                listAux = sectorText[cont].get().split("/")
                techaux = techText[cont].get().upper().split()
                cont2 = 0
                sectorAux = listAux[0].split()
                
                if len(listAux) == 1 and len(sectorAux) == 1:
                    mailText.insert(END," en el sector")
                    textoAux += " en el sector"
                else:
                    mailText.insert(END," en los sectores")
                    textoAux += " en los sectores"
                

                for j in techaux:
                    sectorNumbers = listAux[cont2].split()
                    #print(f"cont2 - {cont2} and {len(techaux)}")
                    
                    cont3 = 0
                    for k in sectorNumbers:
                        mailText.insert(END,f" {j}{k}")
                        textoAux += f" {j}{k}"
                        #print(f"{cont3} - {len(sectorNumbers)} - {sectorNumbers[len(sectorNumbers)-1]} -{k} - {cont2} - {len(techaux)-1}")
                        if cont3 < len(sectorNumbers)   and cont2 != len(techaux)-1:
                            mailText.insert(END,",")
                            textoAux += ","
                        else:
                            if cont3 < len(sectorNumbers)-1:
                                mailText.insert(END,",")
                                textoAux += ","
                            else:
                                mailText.insert(END,".")
                                textoAux += "."
                        cont3 += 1
                    cont2 += 1
            if optionList[cont] == "Sin MIMO Rank2":
                textoAux = ""
                mailText.insert(END,f"{textoSinMIMORank2}")
                textoAux += f"{textoSinMIMORank2}"
                listAux = sectorText[cont].get().split("/")
                techaux = techText[cont].get().upper().split()
                cont2 = 0
                sectorAux = listAux[0].split()
                
                if len(listAux) == 1 and len(sectorAux) == 1:
                    mailText.insert(END," en el sector")
                    textoAux += " en el sector"
                else:
                    mailText.insert(END," en los sectores")
                    textoAux += " en los sectores"
                

                for j in techaux:
                    sectorNumbers = listAux[cont2].split()
                    #print(f"cont2 - {cont2} and {len(techaux)}")
                    
                    cont3 = 0
                    for k in sectorNumbers:
                        mailText.insert(END,f" {j}{k}")
                        textoAux += f" {j}{k}"
                        #print(f"{cont3} - {len(sectorNumbers)} - {sectorNumbers[len(sectorNumbers)-1]} -{k} - {cont2} - {len(techaux)-1}")
                        if cont3 < len(sectorNumbers)   and cont2 != len(techaux)-1:
                            mailText.insert(END,",")
                            textoAux += ","
                        else:
                            if cont3 < len(sectorNumbers)-1:
                                mailText.insert(END,",")
                                textoAux += ","
                            else:
                                mailText.insert(END,".")
                                textoAux += "."
                        cont3 += 1
                    cont2 += 1

            if optionList[cont] == "Sin MIMO Rank4":
                textoAux = ""
                mailText.insert(END,f"{textoSinMIMORank4}")
                textoAux += f"{textoSinMIMORank4}"
                listAux = sectorText[cont].get().split("/")
                techaux = techText[cont].get().upper().split()
                cont2 = 0
                sectorAux = listAux[0].split()
                
                if len(listAux) == 1 and len(sectorAux) == 1:
                    mailText.insert(END," en el sector")
                    textoAux += " en el sector"
                else:
                    mailText.insert(END," en los sectores")
                    textoAux += " en los sectores"
                

                for j in techaux:
                    sectorNumbers = listAux[cont2].split()
                    #print(f"cont2 - {cont2} and {len(techaux)}")
                    
                    cont3 = 0
                    for k in sectorNumbers:
                        mailText.insert(END,f" {j}{k}")
                        textoAux += f" {j}{k}"
                        #print(f"{cont3} - {len(sectorNumbers)} - {sectorNumbers[len(sectorNumbers)-1]} -{k} - {cont2} - {len(techaux)-1}")
                        if cont3 < len(sectorNumbers)   and cont2 != len(techaux)-1:
                            mailText.insert(END,",")
                            textoAux += ","
                        else:
                            if cont3 < len(sectorNumbers)-1:
                                mailText.insert(END,",")
                                textoAux += ","
                            else:
                                mailText.insert(END,".")
                                textoAux += "."
                        cont3 += 1
                    cont2 += 1

            if optionList[cont] == "MIMO Rank2 Bajo":
                textoAux = ""
                mailText.insert(END,f"{textoMIMORank2Bajo}")
                textoAux += f"{textoMIMORank2Bajo}"
                listAux = sectorText[cont].get().split("/")
                techaux = techText[cont].get().upper().split()
                cont2 = 0
                sectorAux = listAux[0].split()
                
                if len(listAux) == 1 and len(sectorAux) == 1:
                    mailText.insert(END," en el sector")
                    textoAux += " en el sector"
                else:
                    mailText.insert(END," en los sectores")
                    textoAux += " en los sectores"
                

                for j in techaux:
                    sectorNumbers = listAux[cont2].split()
                    #print(f"cont2 - {cont2} and {len(techaux)}")
                    
                    cont3 = 0
                    for k in sectorNumbers:
                        mailText.insert(END,f" {j}{k}")
                        textoAux += f" {j}{k}"
                        #print(f"{cont3} - {len(sectorNumbers)} - {sectorNumbers[len(sectorNumbers)-1]} -{k} - {cont2} - {len(techaux)-1}")
                        if cont3 < len(sectorNumbers)   and cont2 != len(techaux)-1:
                            mailText.insert(END,",")
                            textoAux += ","
                        else:
                            if cont3 < len(sectorNumbers)-1:
                                mailText.insert(END,",")
                                textoAux += ","
                            else:
                                mailText.insert(END,".")
                                textoAux += "."
                        cont3 += 1
                    cont2 += 1

            if optionList[cont] == "MIMO Rank4 Bajo":
                textoAux = ""
                mailText.insert(END,f"{textoMIMORank4Bajo}")
                textoAux += f"{textoMIMORank4Bajo}"
                listAux = sectorText[cont].get().split("/")
                techaux = techText[cont].get().upper().split()
                cont2 = 0
                sectorAux = listAux[0].split()
                
                if len(listAux) == 1 and len(sectorAux) == 1:
                    mailText.insert(END," en el sector")
                    textoAux += " en el sector"
                else:
                    mailText.insert(END," en los sectores")
                    textoAux += " en los sectores"
                

                for j in techaux:
                    sectorNumbers = listAux[cont2].split()
                    #print(f"cont2 - {cont2} and {len(techaux)}")
                    
                    cont3 = 0
                    for k in sectorNumbers:
                        mailText.insert(END,f" {j}{k}")
                        textoAux += f" {j}{k}"
                        #print(f"{cont3} - {len(sectorNumbers)} - {sectorNumbers[len(sectorNumbers)-1]} -{k} - {cont2} - {len(techaux)-1}")
                        if cont3 < len(sectorNumbers)   and cont2 != len(techaux)-1:
                            mailText.insert(END,",")
                            textoAux += ","
                        else:
                            if cont3 < len(sectorNumbers)-1:
                                mailText.insert(END,",")
                                textoAux += ","
                            else:
                                mailText.insert(END,".")
                                textoAux += "."
                        cont3 += 1
                    cont2 += 1

            if optionList[cont] == "CA PCELL":
                textoAux = ""
                mailText.insert(END,f"{textoSinCAPCELL}")
                textoAux += f"{textoSinCAPCELL}"
                listAux = sectorText[cont].get().split("/")
                techaux = techText[cont].get().upper().split()
                cont2 = 0
                sectorAux = listAux[0].split()
                
                if len(listAux) == 1 and len(sectorAux) == 1:
                    mailText.insert(END," en el sector")
                    textoAux += " en el sector"
                else:
                    mailText.insert(END," en los sectores")
                    textoAux += " en los sectores"
                

                for j in techaux:
                    sectorNumbers = listAux[cont2].split()
                    #print(f"cont2 - {cont2} and {len(techaux)}")
                    
                    cont3 = 0
                    for k in sectorNumbers:
                        mailText.insert(END,f" {j}{k}")
                        textoAux += f" {j}{k}"
                        #print(f"{cont3} - {len(sectorNumbers)} - {sectorNumbers[len(sectorNumbers)-1]} -{k} - {cont2} - {len(techaux)-1}")
                        if cont3 < len(sectorNumbers)   and cont2 != len(techaux)-1:
                            mailText.insert(END,",")
                            textoAux += ","
                        else:
                            if cont3 < len(sectorNumbers)-1:
                                mailText.insert(END,",")
                                textoAux += ","
                            else:
                                mailText.insert(END,".")
                                textoAux += "."
                        cont3 += 1
                    cont2 += 1
            
            if optionList[cont] == "CA SCELL":
                textoAux = ""
                mailText.insert(END,f"{textoSinCASCELL}")
                textoAux += f"{textoSinCASCELL}"
                listAux = sectorText[cont].get().split("/")
                techaux = techText[cont].get().upper().split()
                cont2 = 0
                sectorAux = listAux[0].split()
                
                if len(listAux) == 1 and len(sectorAux) == 1:
                    mailText.insert(END," en el sector")
                    textoAux += " en el sector"
                else:
                    mailText.insert(END," en los sectores")
                    textoAux += " en los sectores"
                

                for j in techaux:
                    sectorNumbers = listAux[cont2].split()
                    #print(f"cont2 - {cont2} and {len(techaux)}")
                    
                    cont3 = 0
                    for k in sectorNumbers:
                        mailText.insert(END,f" {j}{k}")
                        textoAux += f" {j}{k}"
                        #print(f"{cont3} - {len(sectorNumbers)} - {sectorNumbers[len(sectorNumbers)-1]} -{k} - {cont2} - {len(techaux)-1}")
                        if cont3 < len(sectorNumbers)   and cont2 != len(techaux)-1:
                            mailText.insert(END,",")
                            textoAux += ","
                        else:
                            if cont3 < len(sectorNumbers)-1:
                                mailText.insert(END,",")
                                textoAux += ","
                            else:
                                mailText.insert(END,".")
                                textoAux += "."
                        cont3 += 1
                    cont2 += 1
            
            if optionList[cont] == "SRVCC":
                textoAux = ""
                mailText.insert(END,f"{textoSRVCC}")
                textoAux += f"{textoSRVCC}"
                listAux = sectorText[cont].get().split("/")
                techaux = techText[cont].get().upper().split()
                cont2 = 0
                sectorAux = listAux[0].split()
                
                if len(listAux) == 1 and len(sectorAux) == 1:
                    mailText.insert(END," en el sector")
                    textoAux += " en el sector"
                else:
                    mailText.insert(END," en los sectores")
                    textoAux += " en los sectores"
                

                for j in techaux:
                    sectorNumbers = listAux[cont2].split()
                    #print(f"cont2 - {cont2} and {len(techaux)}")
                    
                    cont3 = 0
                    for k in sectorNumbers:
                        mailText.insert(END,f" {j}{k}")
                        textoAux += f" {j}{k}"
                        #print(f"{cont3} - {len(sectorNumbers)} - {sectorNumbers[len(sectorNumbers)-1]} -{k} - {cont2} - {len(techaux)-1}")
                        if cont3 < len(sectorNumbers)   and cont2 != len(techaux)-1:
                            mailText.insert(END,",")
                            textoAux += ","
                        else:
                            if cont3 < len(sectorNumbers)-1:
                                mailText.insert(END,",")
                                textoAux += ","
                            else:
                                mailText.insert(END,".")
                                textoAux += "."
                        cont3 += 1
                    cont2 += 1
            
            if optionList[cont] == "Sin tráfico 5G":
                textoAux = ""
                mailText.insert(END,f"{textoTrafico5G}")
                textoAux += f"{textoTrafico5G}"
                listAux = sectorText[cont].get().split("/")
                techaux = techText[cont].get().upper().split()
                cont2 = 0
                sectorAux = listAux[0].split()
                
                if len(listAux) == 1 and len(sectorAux) == 1:
                    mailText.insert(END," en el sector")
                    textoAux += " en el sector"
                else:
                    mailText.insert(END," en los sectores")
                    textoAux += " en los sectores"
                

                for j in techaux:
                    sectorNumbers = listAux[cont2].split()
                    #print(f"cont2 - {cont2} and {len(techaux)}")
                    
                    cont3 = 0
                    for k in sectorNumbers:
                        mailText.insert(END,f" {j}{k}")
                        textoAux += f" {j}{k}"
                        #print(f"{cont3} - {len(sectorNumbers)} - {sectorNumbers[len(sectorNumbers)-1]} -{k} - {cont2} - {len(techaux)-1}")
                        if cont3 < len(sectorNumbers)   and cont2 != len(techaux)-1:
                            mailText.insert(END,",")
                            textoAux += ","
                        else:
                            if cont3 < len(sectorNumbers)-1:
                                mailText.insert(END,",")
                                textoAux += ","
                            else:
                                mailText.insert(END,".")
                                textoAux += "."
                        cont3 += 1
                    cont2 += 1

            if optionList[cont] == "PUSCH" or optionList[cont] == "RSSI" or optionList[cont] == "Sin MIMO Rank2" or optionList[cont] == "Sin MIMO Rank4" or optionList[cont] == "MIMO Rank2 Bajo" or optionList[cont] == "MIMO Rank4 Bajo":
                mailText.insert(END," Por favor ¿podrían comprobar que la configuración es correcta o si hay alarmas?\n\n")
                textoAux += " Por favor ¿podrían comprobar que la configuración es correcta o si hay alarmas?"
            if optionList[cont] == "CA PCELL" or optionList[cont] == "CA SCELL" or optionList[cont] == "SRVCC":
                mailText.insert(END," Por favor ¿podrían cargar el script adjunto?\n\n")
                textoAux += " Por favor ¿podrían cargar el script adjunto?"
            if optionList[cont] == "Sin tráfico 5G":
                mailText.insert(END," Por favor ¿podrían comprobar que la configuración X2 es correcta?\n\n")
                textoAux += " Por favor ¿podrían comprobar que la configuración X2 es correcta?"
            
            word.add_paragraph(textoAux)
            if optionList[cont] == "PUSCH":
                try:
                    word.add_picture(".\\pictures\\pusch.png")
                    mailText.insert(END,"\t-> {PUSCH image}\n\n")
                except:
                    messagebox.showinfo("Information","Failure to load image: PUSCH")
            if optionList[cont] == "RSSI":
                try:
                    word.add_picture(".\\pictures\\rssi.png")
                    mailText.insert(END,"\t-> {RSSI image}\n\n")
                except:
                    messagebox.showinfo("Information","Failure to load image: RSSI")
            if optionList[cont] == "Sin MIMO Rank2":
                try:
                    word.add_picture(".\\pictures\\sinmimorank2.png")
                    mailText.insert(END,"\t-> {MIMO Rank2 image}\n\n")
                except:
                    messagebox.showinfo("Information","Failure to load image: Sin MIMO Rank2")
            if optionList[cont] == "Sin MIMO Rank4":
                try:
                    word.add_picture(".\\pictures\\sinmimorank4.png")
                    mailText.insert(END,"\t-> {MIMO Rank4 image}\n\n")
                except:
                    messagebox.showinfo("Information","Failure to load image: Sin MIMO Rank4")
            if optionList[cont] == "MIMO Rank2 Bajo":
                try:
                    word.add_picture(".\\pictures\\mimorank2bajo.png")
                    mailText.insert(END,"\t-> {MIMO Rank2 image}\n\n")
                except:
                    messagebox.showinfo("Information","Failure to load image: MIMO Rank2 Bajo")
            if optionList[cont] == "MIMO Rank4 Bajo":
                try:
                    word.add_picture(".\\pictures\\mimorank4bajo.png")
                    mailText.insert(END,"\t-> {MIMO Rank4 image}\n\n")
                except:
                    messagebox.showinfo("Information","Failure to load image: MIMO Rank4 Bajo")
            if optionList[cont] == "CA PCELL":
                try:
                    word.add_picture(".\\pictures\\capcell.png")
                    mailText.insert(END,"\t-> {CA PCELL image}\n\n")
                except:
                    messagebox.showinfo("Information","Failure to load image: CA PCELL")
            if optionList[cont] == "CA SCELL":
                try:
                    word.add_picture(".\\pictures\\cascell.png")
                    mailText.insert(END,"\t-> {CA SCELL image}\n\n")
                except:
                    messagebox.showinfo("Information","Failure to load image: CA SCELL")
            if optionList[cont] == "SRVCC":
                try:
                    word.add_picture(".\\pictures\\srvcc.png")
                    mailText.insert(END,"\t-> {SRVCC image}\n\n")
                except:
                    messagebox.showinfo("Information","Failure to load image: SRVCC")
            if optionList[cont] == "Sin tráfico 5G":
                try:
                    word.add_picture(".\\pictures\\trafico5g.png")
                    mailText.insert(END,"\t-> {Sin tráfico 5G image}\n\n")
                except:
                    messagebox.showinfo("Information","Failure to load image: Sin tráfico 5G")
                    
        cont += 1
    
def save():
    wordName = siteName.get().upper() + " - email.docx"

    try:
        word.save(wordName)
        messagebox.showinfo("Information",f"Doc <{wordName}> successfully saved")
    except:
        messagebox.showinfo("Information","Something went wrong :(")
    

def clean():
    mailText.delete(1.0,END)
    siteName.set("")
    techTextCA.set("0")
    otherTechTextCA.set("0")
    numberSectorsTextCA.set("0")
    nodeTextCA.set("0")
    enodeTextCA.set("0")
    techTextSRVCC.set("0")
    nodeTextSRVCC.set("0")
    numberSectorsTextSRVCC.set("0")
    for i in techText:
        i.set("0")
    for i in sectorText:
        i.set("0")
    for i in techBox:
        i.config(state=DISABLED)
    for i in sectorBox:
        i.config(state=DISABLED)

    for i in optionsVariables:
        i.set(0)

optionList = ["PUSCH","RSSI","Sin MIMO Rank2","MIMO Rank2 Bajo","Sin MIMO Rank4","MIMO Rank4 Bajo","CA PCELL",
              "CA SCELL","SRVCC","Sin llamadas 4G","Sin tráfico 5G"]

buttonsList = ["Write","Save","Clean"]


# Initiating the app
app = Tk()

# Windows configuration
app.geometry("960x610+0+0")
app.resizable(0,0)
app.title("C44 - Mail generator")
app.config(bg="white")
small_icon = PhotoImage(file=".\\images\\iconC44.png")
big_icon = PhotoImage(file=".\\images\\logoC44.png")
app.iconphoto(False, big_icon, small_icon)

# App TOP
topPanel = Frame(app,bd=2,relief=FLAT)
topPanel.pack(side=TOP)

"""title_name = Label(topPanel,text="C44\t\t\t",fg="black",font=("Dosis",22,"bold"),bg="SlateBlue",width=27) 
title_name.grid(row=0,column=0)"""
logo = PhotoImage(file=".\\images\\connect.png")
title_name = Label(topPanel,image=logo,width=500)
title_name.grid(row=0,column=0)

title_name_2 = Label(topPanel,text="Mail Generator  ",fg="black",font=("Dosis",22,"bold"),bg="cornflowerblue",width=27) 
title_name_2.grid(row=0,column=1)

sitePanel = Frame(topPanel,bd=2,relief=FLAT)
sitePanel.grid(row=1,column=0)

title_name_2 = Label(sitePanel,text="Site Name",fg="black",font=("Dosis",14,"bold"),bg="gray",width=15) 
title_name_2.grid(row=1,column=0)

siteName =StringVar()
siteName.set("")
siteNameBox = Entry(sitePanel,font=("Arial",14),bd=1,width=15,state=NORMAL,textvariable=siteName)
siteNameBox.grid(row=1,column=1)

# App Left
leftPanel = Frame(app,bd=2,relief=FLAT)
leftPanel.pack(side=LEFT)

bottomPanel = Frame(leftPanel,bd=2,relief=FLAT)
bottomPanel.pack(side=BOTTOM)

srvccPanel = LabelFrame(bottomPanel,bd=2,relief=FLAT,bg="gray")
srvccPanel.pack(side=BOTTOM)

caPanel = LabelFrame(bottomPanel,bd=2,relief=FLAT,bg="gray")
caPanel.pack(side=BOTTOM)

optionsPanel = LabelFrame(leftPanel,text="Options\t\tTech     Sectors", font=("Dosis",15,"bold"),bd=1,relief=FLAT,fg="black")
optionsPanel.pack(side=LEFT)


# App Right
rightPanel = Frame(app,bd=2,relief=FLAT)
rightPanel.pack(side=RIGHT)

# Mail Panel
mailPanel = Frame(rightPanel,bd=2,relief=FLAT,bg="black")
mailPanel.pack()

# Buttons Panel
buttonsPanel = Frame(rightPanel,bd =1,relief=FLAT,bg="black")
buttonsPanel.pack()



# Option Panel Configuration

on_image = PhotoImage(width=44, height=20)
off_image = PhotoImage(width=44, height=20)
on_image.put(("green",), to=(0, 0, 23,23))
off_image.put(("red",), to=(24, 0, 47, 23))

optionsVariables= []
techBox = []
techText = []
sectorBox = []
sectorText = []

cont = 0

for option in optionList:
    # Check Buttons
    optionsVariables.append("")
    optionsVariables[cont] = IntVar()
    option = Checkbutton(optionsPanel,image=off_image,selectimage=on_image,indicatoron=False,text=option,font=("Arial",14,"bold"),onvalue=1,offvalue=0, variable=optionsVariables[cont],command=check_button)
    option.grid(row=cont, column=0,sticky=W)  
    
    title_option = Label(optionsPanel,text=optionList[cont],font=("Arial",14,"bold"),width=15) 
    title_option.grid(row=cont,column=1,ipady=3,sticky=W)
    
    # Option Box 
    techBox.append("")
    techText.append("")
    techText[cont] =StringVar()
    techText[cont].set("0")
    techBox[cont] = Entry(optionsPanel,font=("Arial",14,"bold"),bd=1,width=7,state=DISABLED,textvariable=techText[cont])
    techBox[cont].grid(row=cont,column=2,sticky=W)

    sectorBox.append("")
    sectorText.append("")
    sectorText[cont] =StringVar()
    sectorText[cont].set("0")
    sectorBox[cont] = Entry(optionsPanel,font=("Arial",14,"bold"),bd=1,width=15,state=DISABLED,textvariable=sectorText[cont])
    sectorBox[cont].grid(row=cont,column=3,sticky=W)

    cont +=1



# Buttons Panel Configuration
buttonsAux = []
columns = 0

for button in buttonsList:
    button = Button(buttonsPanel,text=button,font=("Arial",13,"bold"),fg="black",bg="Gray",bd=2,width=14)
    buttonsAux.append(button)
    button.grid(row=0,column=columns)
    columns += 1

buttonsAux[0].config(command = write_mail)
buttonsAux[1].config(command = save)
buttonsAux[2].config(command = clean)

# Mail Panel Configuration
mailText = Text(mailPanel,font=("Arial",10),fg="black",bd=3,width=63,height=30)
mailText.grid(row=0,column=0)

# SRVCC Panel Configuration
title_tech_SRVCC = Label(srvccPanel,text="Tech",fg="black",font=("Arial",13),bg="gray",width=5) 
title_tech_SRVCC.grid(row=0,column=1)


title_node_SRVCC = Label(srvccPanel,text="Node",fg="black",font=("Arial",13),bg="gray",width=9) 
title_node_SRVCC.grid(row=0,column=2)

title_node_SRVCC = Label(srvccPanel,text="Nº Cells",fg="black",font=("Arial",13),bg="gray",width=9) 
title_node_SRVCC.grid(row=0,column=3)


empty = Label(srvccPanel,text="",fg="black",font=("Arial",13),bg="gray",width=14) 
empty.grid(row=0,column=4,columnspan=2)

title_name_SRVCC = Label(srvccPanel,text="SRVCC",fg="black",font=("Arial",13,"bold"),bg="gray",width=6) 
title_name_SRVCC.grid(row=1,column=0)

techTextSRVCC =StringVar()
techTextSRVCC.set("0")
techBoxSRVCC = Entry(srvccPanel,font=("Arial",12,),bd=1,width=5,state=NORMAL,textvariable=techTextSRVCC)
techBoxSRVCC.grid(row=1,column=1)


nodeTextSRVCC =StringVar()
nodeTextSRVCC.set("0")
nodeBoxSRVCC = Entry(srvccPanel,font=("Arial",12,),bd=1,width=9,state=NORMAL,textvariable=nodeTextSRVCC)
nodeBoxSRVCC.grid(row=1,column=2)

numberSectorsTextSRVCC =StringVar()
numberSectorsTextSRVCC.set("0")
numberSectorsBoxSRVCC = Entry(srvccPanel,font=("Arial",12,),bd=1,width=9,state=NORMAL,textvariable=numberSectorsTextSRVCC)
numberSectorsBoxSRVCC.grid(row=1,column=3)

# SRVCC Button Panel Configuration

buttonsListSRVCC = "Save"
buttonSRVCC = Button(srvccPanel,text=buttonsListSRVCC,font=("Arial",13,"bold"),fg="black",bg="Gray",bd=2,width=6,command=saveSRVCC)
buttonSRVCC.grid(row=1,column=6)



# CA Panel Configuration
title_tech_CA = Label(caPanel,text="Tech",fg="black",font=("Arial",13),bg="gray",width=5) 
title_tech_CA.grid(row=0,column=1)

title_other_CA = Label(caPanel,text="Other",fg="black",font=("Arial",13),bg="gray",width=9) 
title_other_CA.grid(row=0,column=2)

title_node_CA = Label(caPanel,text="Node",fg="black",font=("Arial",13),bg="gray",width=9) 
title_node_CA.grid(row=0,column=3)

title_node_CA = Label(caPanel,text="EnodeB",fg="black",font=("Arial",13),bg="gray",width=6) 
title_node_CA.grid(row=0,column=4)

title_node_CA = Label(caPanel,text="Nº Cells",fg="black",font=("Arial",13),bg="gray",width=7) 
title_node_CA.grid(row=0,column=5)

title_name_CA = Label(caPanel,text="CA",fg="black",font=("Arial",13,"bold"),bg="gray",width=6) 
title_name_CA.grid(row=1,column=0)

techTextCA =StringVar()
techTextCA.set("0")
techBoxCA = Entry(caPanel,font=("Arial",12,),bd=1,width=5,state=NORMAL,textvariable=techTextCA)
techBoxCA.grid(row=1,column=1)

otherTechTextCA =StringVar()
otherTechTextCA.set("0")
otherTechBoxCA = Entry(caPanel,font=("Arial",12,),bd=1,width=9,state=NORMAL,textvariable=otherTechTextCA)
otherTechBoxCA.grid(row=1,column=2)

nodeTextCA =StringVar()
nodeTextCA.set("0")
nodeBoxCA = Entry(caPanel,font=("Arial",12,),bd=1,width=9,state=NORMAL,textvariable=nodeTextCA)
nodeBoxCA.grid(row=1,column=3)

enodeTextCA =StringVar()
enodeTextCA.set("0")
enodeBoxCA = Entry(caPanel,font=("Arial",12,),bd=1,width=6,state=NORMAL,textvariable=enodeTextCA)
enodeBoxCA.grid(row=1,column=4)

numberSectorsTextCA =StringVar()
numberSectorsTextCA.set("0")
numberSectorsBoxCA = Entry(caPanel,font=("Arial",12,),bd=1,width=7,state=NORMAL,textvariable=numberSectorsTextCA)
numberSectorsBoxCA.grid(row=1,column=5)

# CA Button Panel Configuration

buttonsListCA = "Save"
button = Button(caPanel,text=buttonsListCA,font=("Arial",13,"bold"),fg="black",bg="Gray",bd=2,width=6,command=saveCA)
button.grid(row=1,column=6)

app.mainloop()