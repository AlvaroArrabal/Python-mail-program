from tkinter import *
from tkinter import messagebox
from modules import SRVCC,CA,justification
import time
from docx import Document

word = Document()

def check_button():
    cont=0
    for i in techBox:
        if optionsVariables[cont].get() == 1:
            techBox[cont].config(state=NORMAL)
            if techBox[cont].get() == "0":
                techBox[cont].delete(0,END)
                techBox[cont].focus()
            
            if optionList[cont] != "Sin datos 5G":
                sectorBox[cont].config(state=NORMAL)
                if sectorBox[cont].get() == "0":
                    sectorBox[cont].delete(0,END)
                    sectorBox[cont].focus()
                
        else:
            techBox[cont].config(state=DISABLED)
            techText[cont].set("0")
            sectorBox[cont].config(state=DISABLED)
            sectorText[cont].set("0")
        cont += 1

def saveSRVCC():
    site = siteName.get().upper()
    tech = techTextSRVCC.get().upper()
    node = nodeTextSRVCC.get()
    cells = numberSectorsTextSRVCC.get()

    try:
        if tech == "0" or node == "0" or cells == "0":
            messagebox.showinfo("Information","Error: Not enough information")
        else:
            SRVCC.crear_txt(site,tech,node,int(cells))
            messagebox.showinfo("Information","Completed!")
    except:
        messagebox.showinfo("Information","Something went wrong :(")

def saveCA():
    try:
        
        site = siteName.get().upper()
        techCA = techTextCA.get().upper()
        otherTechs = otherTechTextCA.get().upper()
        cells = numberSectorsTextCA.get()
        node = nodeTextCA.get().upper()
        enodeB = enodeTextCA.get().upper()
        if site == "":
            messagebox.showinfo("Information","Error: Empty site name")
        elif techCA == "0" or otherTechs == "0" or cells == "0" or node == "0" or enodeB == "0":
            messagebox.showinfo("Information","Error: Not enough information")
        elif techCA == "Y" and "M" in otherTechs and len(otherTechs) == 1:
            messagebox.showinfo("Information","Error: Not necessary CA between Y and M.")
        else:
            CA.crear_txt_CA(site,techCA,otherTechs,int(cells),node,enodeB)
            messagebox.showinfo("Information","Completed!")
    except:
        messagebox.showinfo("Information","Something went wrong :(")

def insert_photo(name):
    name.lower()
    try:
        word.add_picture(f".\\pictures\\{name}.png")
        mailText.insert(END,f"\t-> {name} image\n\n")
    except:
        messagebox.showinfo("Information",f"Failure to load image: {name}")

def write_mail():
    mailText.delete(1.0 ,END)
    word._body.clear_content()
    now = time.strftime("%X")
    
    if now < str(12):
        mailText.insert(END,"Buenos días,\n")
        word.add_paragraph("Buenos días,")
    else:
        mailText.insert(END,"Buenas tardes,\n")
        word.add_paragraph("Buenas tardes,")

    mailText.insert(END,f"Hemos detectado las siguientes inconsistencias en el site {siteName.get().upper()}:\n")
    word.add_paragraph(f"Hemos detectado las siguientes inconsistencias en el site {siteName.get().upper()}:")

    cont = 0
    for i in techText:
        if i.get() != "0":
            
            textoAux = justification.insert(sectorText[cont],techText[cont],optionList[cont])
            mailText.insert(END,textoAux)
            word.add_paragraph(textoAux)
            #insert_photo(optionList[cont])
        cont += 1
    
def save():
    wordName = siteName.get().upper() + " - email.docx"

    try:
        word.save(wordName)
        messagebox.showinfo("Information",f"Doc <{wordName}> successfully saved")
    except:
        messagebox.showinfo("Information","Something went wrong :(")
    

def clean():
    mailText.delete(1.0,END)
    word._body.clear_content()
    siteName.set("")
    techTextCA.set("0")
    otherTechTextCA.set("0")
    numberSectorsTextCA.set("0")
    nodeTextCA.set("0")
    enodeTextCA.set("0")
    techTextSRVCC.set("0")
    nodeTextSRVCC.set("0")
    numberSectorsTextSRVCC.set("0")
    for i in techText:
        i.set("0")
    for i in sectorText:
        i.set("0")
    for i in techBox:
        i.config(state=DISABLED)
    for i in sectorBox:
        i.config(state=DISABLED)

    for i in optionsVariables:
        i.set(0)

optionList =[  
            "ICM Band",
            "RTWP", 
            "PUSCH",
            "RSSI",
            "Speech Discon.(3dB)",
            "3G Ending 2G (3dB)",
            "Sin MIMO Rank2",
            "MIMO Rank2 Bajo",
            "Sin MIMO Rank4",
            "MIMO Rank4 Bajo",
            "CSFallBack",
            "CA PCELL",
            "CA SCELL",
            "SRVCC",
            "Sin llamadas 2G",
            "Sin llamadas 3G",
            "Sin llamadas 4G",
            "Sin tráfico 5G" ,
            "Sin datos 5G"
            ]

buttonsList = ["Write","Save","Clean"]

# Initiating the app
app = Tk()

# Windows configuration
app.geometry("950x680+0+0")
app.resizable(0,0)
app.title("C44 - Mail generator")
app.config(bg="WhiteSmoke")
small_icon = PhotoImage(file=".\\images\\iconC44.png")
big_icon = PhotoImage(file=".\\images\\logoC44.png")
app.iconphoto(False, big_icon, small_icon)

# App TOP
topPanel = Frame(app,bd=2,relief=FLAT)
topPanel.pack(side=TOP)

logo = PhotoImage(file=".\\images\\connect.png")
title_name = Label(topPanel,image=logo,width=500)
title_name.grid(row=0,column=0)

title_name_2 = Label(topPanel,text="Mail Generator  ",fg="black",font=("Dosis",22,"bold"),bg="cornflowerblue",width=27) 
title_name_2.grid(row=0,column=1)

sitePanel = Frame(topPanel,bd=2,relief=FLAT)
sitePanel.grid(row=1,column=0)

title_name_2 = Label(sitePanel,text="Site Name",fg="black",font=("Dosis",14,"bold"),bg="gray",width=15) 
title_name_2.grid(row=1,column=0)

siteName =StringVar()
siteName.set("")
siteNameBox = Entry(sitePanel,font=("Arial",14),bd=1,width=15,state=NORMAL,textvariable=siteName)
siteNameBox.grid(row=1,column=1)

# App Left
leftPanel = Frame(app,bd=2,relief=FLAT)
leftPanel.pack(side=LEFT)

bottomPanel = Frame(leftPanel,bd=2,relief=FLAT)
bottomPanel.pack(side=BOTTOM)



optionsPanel = LabelFrame(leftPanel,text="\tOptions\t  Tech\tSectors", font=("Dosis",15,"bold"),bd=1,relief=FLAT,fg="black")
optionsPanel.pack(side=LEFT)


# App Right
rightPanel = Frame(app,bd=2,relief=FLAT)
rightPanel.pack(side=RIGHT)

srvccPanel = LabelFrame(rightPanel,bd=2,relief=FLAT,bg="gray")
srvccPanel.pack(side=BOTTOM)

caPanel = LabelFrame(rightPanel,bd=2,relief=FLAT,bg="gray")
caPanel.pack(side=BOTTOM)

# Mail Panel
mailPanel = Frame(rightPanel,bd=2,relief=FLAT,bg="black")
mailPanel.pack()

# Buttons Panel
buttonsPanel = Frame(rightPanel,bd =1,relief=FLAT,bg="black")
buttonsPanel.pack()



# Option Panel Configuration

on_image = PhotoImage(width=44, height=20)
off_image = PhotoImage(width=44, height=20)
on_image.put(("darkgreen",), to=(0, 0, 22,23))
off_image.put(("crimson",), to=(24, 0, 47, 24))


optionsVariables= []
techBox = []
techText = []
sectorBox = []
sectorText = []

cont = 0

for option in optionList:
    # Check Buttons
    optionsVariables.append("")
    optionsVariables[cont] = IntVar()
    option = Checkbutton(optionsPanel,image=off_image,selectimage=on_image,indicatoron=False,text=option,font=("Arial",11,"bold"),onvalue=1,offvalue=0, variable=optionsVariables[cont],command=check_button)
    option.grid(row=cont, column=0,sticky=W)  
    
    title_option = Label(optionsPanel,text=optionList[cont],font=("Arial",11,"bold"),width=15) 
    title_option.grid(row=cont,column=1,ipady=3,sticky=W)
    
    # Option Box 
    techBox.append("")
    techText.append("")
    techText[cont] =StringVar()
    techText[cont].set("0")
    techBox[cont] = Entry(optionsPanel,font=("Arial",14,"bold"),bd=1,width=7,state=DISABLED,textvariable=techText[cont])
    techBox[cont].grid(row=cont,column=2,sticky=W)

    sectorBox.append("")
    sectorText.append("")
    sectorText[cont] =StringVar()
    
    if optionList[cont] == "Sin datos 5G":
        sectorText[cont].set("")
    else:
        sectorText[cont].set("0")
    sectorBox[cont] = Entry(optionsPanel,font=("Arial",14,"bold"),bd=1,width=15,state=DISABLED,textvariable=sectorText[cont])
    sectorBox[cont].grid(row=cont,column=3,sticky=W)

    cont +=1



# Buttons Panel Configuration
buttonsAux = []
columns = 0

for button in buttonsList:
    button = Button(buttonsPanel,text=button,font=("Arial",13,"bold"),fg="black",bg="Gray",bd=2,width=15)
    buttonsAux.append(button)
    button.grid(row=0,column=columns)
    columns += 1

buttonsAux[0].config(command = write_mail)
buttonsAux[1].config(command = save)
buttonsAux[2].config(command = clean)

# Mail Panel Configuration
mailText = Text(mailPanel,font=("Arial",10),fg="black",bd=3,width=66,height=26)
mailText.grid(row=0,column=0)

# SRVCC Panel Configuration
title_tech_SRVCC = Label(srvccPanel,text="Tech",fg="black",font=("Arial",13),bg="gray",width=5) 
title_tech_SRVCC.grid(row=0,column=1)


title_node_SRVCC = Label(srvccPanel,text="Node",fg="black",font=("Arial",13),bg="gray",width=9) 
title_node_SRVCC.grid(row=0,column=2)

title_node_SRVCC = Label(srvccPanel,text="Nº Cells",fg="black",font=("Arial",13),bg="gray",width=9) 
title_node_SRVCC.grid(row=0,column=3)


empty = Label(srvccPanel,text="",fg="black",font=("Arial",13),bg="gray",width=14) 
empty.grid(row=0,column=4,columnspan=2)

title_name_SRVCC = Label(srvccPanel,text="SRVCC",fg="black",font=("Arial",13,"bold"),bg="gray",width=6) 
title_name_SRVCC.grid(row=1,column=0)

techTextSRVCC =StringVar()
techTextSRVCC.set("0")
techBoxSRVCC = Entry(srvccPanel,font=("Arial",12,),bd=1,width=5,state=NORMAL,textvariable=techTextSRVCC)
techBoxSRVCC.grid(row=1,column=1)


nodeTextSRVCC =StringVar()
nodeTextSRVCC.set("0")
nodeBoxSRVCC = Entry(srvccPanel,font=("Arial",12,),bd=1,width=9,state=NORMAL,textvariable=nodeTextSRVCC)
nodeBoxSRVCC.grid(row=1,column=2)

optionSector =["1","2","3","4"]

numberSectorsTextSRVCC =StringVar()
numberSectorsTextSRVCC.set("0")
numberSectorsBoxSRVCC = OptionMenu(srvccPanel,numberSectorsTextSRVCC,*optionSector)
numberSectorsBoxSRVCC.grid(row=1,column=3)

# SRVCC Button Panel Configuration

buttonsListSRVCC = "Save"
buttonSRVCC = Button(srvccPanel,text=buttonsListSRVCC,font=("Arial",13,"bold"),fg="black",bg="Gray",bd=2,width=6,command=saveSRVCC)
buttonSRVCC.grid(row=1,column=6)



# CA Panel Configuration
title_tech_CA = Label(caPanel,text="Tech",fg="black",font=("Arial",13),bg="gray",width=5) 
title_tech_CA.grid(row=0,column=1)

title_other_CA = Label(caPanel,text="Other",fg="black",font=("Arial",13),bg="gray",width=9) 
title_other_CA.grid(row=0,column=2)

title_node_CA = Label(caPanel,text="Node",fg="black",font=("Arial",13),bg="gray",width=9) 
title_node_CA.grid(row=0,column=3)

title_node_CA = Label(caPanel,text="EnodeB",fg="black",font=("Arial",13),bg="gray",width=6) 
title_node_CA.grid(row=0,column=4)

title_node_CA = Label(caPanel,text="Nº Cells",fg="black",font=("Arial",13),bg="gray",width=7) 
title_node_CA.grid(row=0,column=5)

title_name_CA = Label(caPanel,text="CA",fg="black",font=("Arial",13,"bold"),bg="gray",width=6) 
title_name_CA.grid(row=1,column=0)

techTextCA =StringVar()
techTextCA.set("0")
techBoxCA = Entry(caPanel,font=("Arial",12,),bd=1,width=5,state=NORMAL,textvariable=techTextCA)
techBoxCA.grid(row=1,column=1)

otherTechTextCA =StringVar()
otherTechTextCA.set("0")
otherTechBoxCA = Entry(caPanel,font=("Arial",12,),bd=1,width=9,state=NORMAL,textvariable=otherTechTextCA)
otherTechBoxCA.grid(row=1,column=2)

nodeTextCA =StringVar()
nodeTextCA.set("0")
nodeBoxCA = Entry(caPanel,font=("Arial",12,),bd=1,width=9,state=NORMAL,textvariable=nodeTextCA)
nodeBoxCA.grid(row=1,column=3)

enodeTextCA =StringVar()
enodeTextCA.set("0")
enodeBoxCA = Entry(caPanel,font=("Arial",12,),bd=1,width=6,state=NORMAL,textvariable=enodeTextCA)
enodeBoxCA.grid(row=1,column=4)

numberSectorsTextCA =StringVar()
numberSectorsTextCA.set("0")
numberSectorsBoxCA = OptionMenu(caPanel,numberSectorsTextCA,*optionSector)
numberSectorsBoxCA.grid(row=1,column=5)

# CA Button Panel Configuration

buttonsListCA = "Save"
button = Button(caPanel,text=buttonsListCA,font=("Arial",13,"bold"),fg="black",bg="Gray",bd=2,width=6,command=saveCA)
button.grid(row=1,column=6)

app.mainloop()